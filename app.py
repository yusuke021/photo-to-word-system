"""
写真貼り付けシステム - Streamlit Webアプリ版
表形式で写真をWordファイルに貼り付ける
"""

import streamlit as st
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os
import io
import tempfile
from datetime import datetime

# HEIC画像のサポート
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    pass

# ページ設定
st.set_page_config(
    page_title="写真貼り付けシステム",
    page_icon="📸",
    layout="wide"
)

# カスタムCSS
st.markdown("""
<style>
    .main {
        background-color: #f5f6fa;
    }
    .stButton>button {
        width: 100%;
        background-color: #3498db;
        color: white;
        font-weight: bold;
        border-radius: 5px;
        padding: 10px;
    }
    .success-box {
        padding: 20px;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        color: #155724;
    }
</style>
""", unsafe_allow_html=True)


def parse_filename(filepath):
    """
    ファイル名を解析して部品名と写真区分を取得
    フォーマット: 番号_部品名_重量_単位_素材ID_加工ID_実施者ID_写真区分_特記事項.拡張子
    """
    try:
        filename = os.path.basename(filepath)
        name_without_ext = os.path.splitext(filename)[0]
        parts = name_without_ext.split('_')
        
        if len(parts) >= 8:
            part_name = parts[1]
            photo_type = parts[7]
            return (part_name, photo_type)
        else:
            return (None, None)
    except Exception:
        return (None, None)


def filter_images_by_photo_type(uploaded_files, insert_name):
    """写真区分がPの画像のみをフィルタリング"""
    if not insert_name:
        return uploaded_files, 0
    
    filtered_images = []
    skipped_count = 0
    
    for uploaded_file in uploaded_files:
        part_name, photo_type = parse_filename(uploaded_file.name)
        
        if part_name is None or photo_type is None:
            skipped_count += 1
            continue
        
        if photo_type.upper() == 'P':
            filtered_images.append(uploaded_file)
        else:
            skipped_count += 1
    
    return filtered_images, skipped_count


def set_table_borders(table, border_type):
    """表の罫線を設定"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    tblBorders = OxmlElement('w:tblBorders')
    
    if border_type == "なし":
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
    elif border_type == "すべて":
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
    elif border_type == "外枠のみ":
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        for border_name in ['insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
    
    tblPr.append(tblBorders)


def insert_image_to_cell(cell, uploaded_file, cell_height_mm, ppi=220):
    """セルに画像を挿入（指定されたPPIで）"""
    img = Image.open(uploaded_file)
    img_width, img_height = img.size
    aspect_ratio = img_width / img_height
    target_height_mm = cell_height_mm
    
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    
    run = paragraph.add_run()
    
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    # 画像を指定のPPIでメモリ上に保存
    if file_ext in ['.heic', '.heif']:
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        
        img_byte_arr = io.BytesIO()
        # DPI情報を設定してJPEG保存
        img.save(img_byte_arr, format='JPEG', quality=95, dpi=(ppi, ppi))
        img_byte_arr.seek(0)
        run.add_picture(img_byte_arr, height=Mm(target_height_mm))
    else:
        # 通常の画像もPPI情報を設定
        uploaded_file.seek(0)
        img_byte_arr = io.BytesIO()
        
        # RGBAモードの場合はRGBに変換
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        
        # 元の形式を保持しつつDPI情報を設定
        img_format = 'JPEG' if file_ext in ['.jpg', '.jpeg'] else 'PNG'
        if img_format == 'JPEG':
            img.save(img_byte_arr, format=img_format, quality=95, dpi=(ppi, ppi))
        else:
            img.save(img_byte_arr, format=img_format, dpi=(ppi, ppi))
        
        img_byte_arr.seek(0)
        run.add_picture(img_byte_arr, height=Mm(target_height_mm))


def insert_part_name_to_cell(cell, part_name):
    """偶数行のセルに部品名を挿入"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run(part_name)
    
    char_count = len(part_name)
    if char_count <= 18:
        font_size = 12
    elif char_count <= 20:
        font_size = 11
    elif char_count <= 22:
        font_size = 10
    elif char_count <= 24:
        font_size = 9
    elif char_count <= 26:
        font_size = 8
    elif char_count <= 28:
        font_size = 7
    else:
        font_size = 7
    
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:eastAsia'), 'MS Mincho')
    r_fonts.set(qn('w:ascii'), 'Times New Roman')
    r_fonts.set(qn('w:hAnsi'), 'Times New Roman')
    r_pr.append(r_fonts)


def create_word_document(uploaded_files, settings, insert_name, existing_doc_file=None, ppi=220):
    """Wordドキュメントを作成または既存ファイルに追記"""
    # 既存のWordファイルがある場合はそれを開く、ない場合は新規作成（テンプレート使用）
    if existing_doc_file is not None:
        doc = Document(existing_doc_file)
    else:
        # テンプレートファイルが存在する場合はそれを使用
        # Streamlit Cloud対応：複数のパスを試行
        template_paths = [
            'template.docx',  # カレントディレクトリ
            os.path.join(os.getcwd(), 'template.docx'),  # 絶対パス
        ]
        
        # __file__が利用可能な場合は追加
        try:
            if '__file__' in globals():
                template_paths.insert(0, os.path.join(os.path.dirname(__file__), 'template.docx'))
        except:
            pass
        
        template_found = False
        for template_path in template_paths:
            if os.path.exists(template_path):
                try:
                    doc = Document(template_path)
                    st.info("📋 テンプレートを使用して新規ファイルを作成します")
                    template_found = True
                    break
                except Exception as e:
                    continue
        
        if not template_found:
            doc = Document()
            st.info("📄 空白のWordファイルを作成します")
    
    rows = settings['rows']
    cols = settings['cols']
    odd_height = settings['odd_height']
    odd_width = settings['odd_width']
    even_height = settings['even_height']
    even_width = settings['even_width']
    border_type = settings['border_type']
    table_align = settings['table_align']
    
    images_per_page = (rows // 2) * cols
    total_images = len(uploaded_files)
    num_pages = (total_images + images_per_page - 1) // images_per_page
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for page_idx in range(num_pages):
        start_idx = page_idx * images_per_page
        end_idx = min(start_idx + images_per_page, total_images)
        page_images = uploaded_files[start_idx:end_idx]
        
        status_text.text(f"ページ {page_idx + 1} / {num_pages} を処理中...")
        progress_bar.progress((page_idx + 1) / num_pages)
        
        table = doc.add_table(rows=rows, cols=cols)
        
        if table_align == "中央":
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif table_align == "右揃え":
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        else:
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for row_idx in range(rows):
            for col_idx in range(cols):
                cell = table.rows[row_idx].cells[col_idx]
                
                if (row_idx + 1) % 2 == 1:
                    cell.width = Mm(odd_width)
                    table.rows[row_idx].height = Mm(odd_height)
                else:
                    cell.width = Mm(even_width)
                    table.rows[row_idx].height = Mm(even_height)
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        set_table_borders(table, border_type)
        
        image_idx = 0
        for row_idx in range(rows):
            if (row_idx + 1) % 2 == 1:
                for col_idx in range(cols):
                    if image_idx < len(page_images):
                        cell = table.rows[row_idx].cells[col_idx]
                        image_file = page_images[image_idx]
                        
                        insert_image_to_cell(cell, image_file, odd_height, ppi)
                        
                        if insert_name and row_idx + 1 < rows:
                            part_name, _ = parse_filename(image_file.name)
                            if part_name:
                                even_row_cell = table.rows[row_idx + 1].cells[col_idx]
                                insert_part_name_to_cell(even_row_cell, part_name)
                        
                        image_idx += 1
        
        if page_idx < num_pages - 1:
            doc.add_page_break()
    
    progress_bar.empty()
    status_text.empty()
    
    return doc


# メインUI
st.title("📸 写真を表形式でWordに貼り付ける")
st.markdown("---")

# Wordファイルアップロード（オプション）
st.header("📄 Wordファイル（オプション）")
uploaded_word = st.file_uploader(
    "既存のWordファイルをアップロード（省略すると新規作成）",
    type=['docx'],
    help="既存のWordファイルに追記したい場合はアップロードしてください。省略した場合、template.docxがあれば自動的にテンプレートが適用されます。"
)

if uploaded_word:
    st.success(f"✅ {uploaded_word.name} が選択されています（このファイルに追記されます）")
else:
    # テンプレートファイルの存在確認
    template_exists = False
    template_paths = ['template.docx', os.path.join(os.getcwd(), 'template.docx')]
    
    # __file__が利用可能な場合は追加
    try:
        if '__file__' in globals():
            template_paths.insert(0, os.path.join(os.path.dirname(__file__), 'template.docx'))
    except:
        pass
    
    for template_path in template_paths:
        if os.path.exists(template_path):
            template_exists = True
            break
    
    if template_exists:
        st.info("📋 テンプレートファイル (template.docx) を使用して新規作成されます")
    else:
        st.info("新規Wordファイルが作成されます（テンプレートなし）")

st.markdown("---")

# サイドバー設定
with st.sidebar:
    st.header("📝 名称挿入設定")
    insert_name = st.checkbox(
        "写真の部品名を挿入",
        value=False,
        help="ファイル名形式: 番号_部品名_..._写真区分_...\n写真区分=Pのみ貼り付け"
    )
    
    st.markdown("---")
    st.header("⚙️ 表の設定")
    
    rows = st.number_input("行数", min_value=2, max_value=20, value=8, step=2)
    cols = st.number_input("列数", min_value=1, max_value=10, value=2, step=1)
    
    st.subheader("📊 罫線設定")
    border_type = st.radio("罫線", ["なし", "すべて", "外枠のみ"], index=0)
    
    st.subheader("📍 表の配置")
    table_align = st.radio("配置", ["左揃え", "中央", "右揃え"], index=1)
    
    st.markdown("---")
    st.header("📐 セルサイズ (mm)")
    
    st.subheader("📷 奇数行（写真用）")
    odd_height = st.number_input("高さ (mm)", min_value=10.0, max_value=200.0, value=50.0, step=1.0, key="odd_h")
    odd_width = st.number_input("幅 (mm)", min_value=10.0, max_value=200.0, value=82.0, step=1.0, key="odd_w")
    
    st.subheader("📝 偶数行（説明用）")
    even_height = st.number_input("高さ (mm)", min_value=5.0, max_value=100.0, value=7.0, step=1.0, key="even_h")
    even_width = st.number_input("幅 (mm)", min_value=10.0, max_value=200.0, value=82.0, step=1.0, key="even_w")
    
    st.markdown("---")
    st.header("🎨 画像品質設定")
    image_quality = st.radio(
        "画像の解像度 (PPI)",
        options=["印刷用 (220 ppi)", "高性能 (300 ppi)", "標準 (150 ppi)"],
        index=0,
        help="印刷用がファイルサイズと品質のバランスが良く推奨されます"
    )
    
    # PPI値を抽出
    ppi_map = {
        "印刷用 (220 ppi)": 220,
        "高性能 (300 ppi)": 300,
        "標準 (150 ppi)": 150
    }
    selected_ppi = ppi_map[image_quality]

# メインコンテンツ
col1, col2 = st.columns([2, 1])

with col1:
    st.header("🖼️ 画像ファイルをアップロード")
    uploaded_files = st.file_uploader(
        "画像を選択してください（複数選択可）",
        type=['jpg', 'jpeg', 'png', 'gif', 'bmp', 'heic', 'HEIC'],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)}枚の画像が選択されています")
        
        # プレビュー表示
        with st.expander("📋 選択された画像リスト"):
            for idx, file in enumerate(uploaded_files, 1):
                part_name, photo_type = parse_filename(file.name)
                if part_name:
                    st.text(f"{idx}. {file.name} - 部品名: {part_name}, 区分: {photo_type}")
                else:
                    st.text(f"{idx}. {file.name}")

with col2:
    st.header("📊 設定サマリー")
    
    # Wordファイルの状態
    if uploaded_word:
        word_status = f"📄 既存ファイル: {uploaded_word.name}"
    else:
        word_status = "📄 新規Wordファイル"
    
    st.info(f"""
    {word_status}
    
    **表の設定:**
    - 行数: {rows}行
    - 列数: {cols}列
    - 罫線: {border_type}
    - 配置: {table_align}
    
    **セルサイズ:**
    - 奇数行: {odd_height}mm × {odd_width}mm
    - 偶数行: {even_height}mm × {even_width}mm
    
    **画像品質:**
    - 解像度: {image_quality}
    
    **1ページあたり:** {(rows // 2) * cols}枚の写真
    """)

st.markdown("---")

# 実行ボタン
if st.button("✨ Wordファイルを生成", type="primary"):
    if not uploaded_files:
        st.error("❌ 画像ファイルを選択してください")
    else:
        with st.spinner("処理中..."):
            # フィルタリング
            filtered_files, skipped = filter_images_by_photo_type(uploaded_files, insert_name)
            
            if skipped > 0:
                st.warning(f"⚠️ {skipped}枚の画像がスキップされました（写真区分がP以外、または規則に従っていない）")
            
            if not filtered_files:
                st.error("❌ 貼り付け可能な画像がありません")
            else:
                settings = {
                    'rows': rows,
                    'cols': cols,
                    'odd_height': odd_height,
                    'odd_width': odd_width,
                    'even_height': even_height,
                    'even_width': even_width,
                    'border_type': border_type,
                    'table_align': table_align
                }
                
                try:
                    doc = create_word_document(filtered_files, settings, insert_name, uploaded_word, selected_ppi)
                    
                    # メモリ上に保存
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    # ダウンロードボタン
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    if uploaded_word:
                        # 既存ファイル名をベースにする
                        base_name = os.path.splitext(uploaded_word.name)[0]
                        filename = f"{base_name}_追記_{timestamp}.docx"
                    else:
                        filename = f"写真貼り付け_{timestamp}.docx"
                    
                    if uploaded_word:
                        st.success(f"✅ 既存のWordファイルに{len(filtered_files)}枚の画像を追記しました！")
                    else:
                        st.success(f"✅ {len(filtered_files)}枚の画像を表に貼り付けました！")
                    
                    st.download_button(
                        label="📥 Wordファイルをダウンロード",
                        data=doc_io,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"❌ エラーが発生しました: {str(e)}")

# フッター
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #888;'>
    <p>写真貼り付けシステム v2.0 | Powered by Streamlit</p>
</div>
""", unsafe_allow_html=True)
